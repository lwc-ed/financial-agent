"""
產生六個模型比較報告 Word 檔（全動態版）
輸出：ml/report/model_comparison_report.docx
"""

from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json
import csv

# ───────────────────────────────────────────────
# 路徑設定
# ───────────────────────────────────────────────
ML_ROOT   = Path(__file__).resolve().parent.parent
OUTPUT_DIR = Path(__file__).resolve().parent
OUTPUT_PATH = OUTPUT_DIR / "model_comparison_report.docx"

MODELS = [
    "bigru",
    "bigru_TL_alignment",
    "bilstm",
    "bilstm_TL_alignment",
    "gru_TL_alignment",
    "xgboost_TL_alignment",
]

MODEL_DISPLAY = {
    "bigru":                "BiGRU",
    "bigru_TL_alignment":   "BiGRU + TL",
    "bilstm":               "BiLSTM",
    "bilstm_TL_alignment":  "BiLSTM + TL",
    "gru_TL_alignment":     "GRU + TL",
    "xgboost_TL_alignment": "XGBoost + TL",
}

# ───────────────────────────────────────────────
# 讀取 metrics
# ───────────────────────────────────────────────
def load_metrics(model_name):
    base = ML_ROOT / "model_outputs" / model_name
    reg = json.loads((base / "metrics_regression.json").read_text())
    alm = json.loads((base / "metrics_alarm_binary.json").read_text())
    r4  = json.loads((base / "metrics_risk_4class.json").read_text())
    return reg, alm, r4

all_metrics = {m: load_metrics(m) for m in MODELS}

# ───────────────────────────────────────────────
# 讀取 predictions.csv 統計（動態）
# ───────────────────────────────────────────────
def load_pred_stats(model_name):
    path = ML_ROOT / "model_outputs" / model_name / "predictions.csv"
    total = 0
    inf_count = 0
    users = set()
    with open(path, encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            total += 1
            users.add(row["user_id"])
            if row.get("true_risk_ratio", "") == "inf":
                inf_count += 1
    return {"total": total, "inf": inf_count, "users": sorted(users)}

pred_stats = {m: load_pred_stats(m) for m in MODELS}

# ───────────────────────────────────────────────
# 動態計算排名
# ───────────────────────────────────────────────
def compute_rankings():
    """
    按 Binary Alarm F1 排名，同分比 MAE（低優先）。
    回傳 list of dict，依排名排序。
    """
    rows = []
    for m in MODELS:
        reg, alm, r4 = all_metrics[m]
        rows.append({
            "model":   m,
            "display": MODEL_DISPLAY[m],
            "mae":     reg["MAE"],
            "rmse":    reg["RMSE"],
            "f1":      alm["F1-score"],
            "macro_f1": r4["Macro F1"],
        })
    rows.sort(key=lambda x: (-x["f1"], x["mae"]))
    for i, r in enumerate(rows):
        r["rank"] = i + 1
    return rows

rankings = compute_rankings()

# ───────────────────────────────────────────────
# 動態生成觀察文字
# ───────────────────────────────────────────────
def obs_regression():
    best_mae_m  = min(MODELS, key=lambda m: all_metrics[m][0]["MAE"])
    best_rmse_m = min(MODELS, key=lambda m: all_metrics[m][0]["RMSE"])
    best_mape_m = min(MODELS, key=lambda m: all_metrics[m][0]["MAPE"])
    best_mae_v  = all_metrics[best_mae_m][0]["MAE"]
    best_rmse_v = all_metrics[best_rmse_m][0]["RMSE"]
    best_mape_v = all_metrics[best_mape_m][0]["MAPE"]
    worst_mae_m = max(MODELS, key=lambda m: all_metrics[m][0]["MAE"])
    worst_mae_v = all_metrics[worst_mae_m][0]["MAE"]
    return (
        f"觀察：{MODEL_DISPLAY[best_mae_m]} 在 MAE 表現最佳（{best_mae_v:.2f}），"
        f"{MODEL_DISPLAY[best_rmse_m]} 在 RMSE 最低（{best_rmse_v:.2f}）。"
        f"{MODEL_DISPLAY[best_mape_m]} 的 MAPE 最低（{best_mape_v:.2f}%）。"
        f"{MODEL_DISPLAY[worst_mae_m]} 的 MAE 最高（{worst_mae_v:.2f}），"
        f"若測試集包含高消費用戶，絕對誤差自然偏大，需留意跨模型比較時的資料集差異。"
    )

def obs_alarm():
    best_f1_m  = max(MODELS, key=lambda m: all_metrics[m][1]["F1-score"])
    worst_f1_m = min(MODELS, key=lambda m: all_metrics[m][1]["F1-score"])
    best_f1_v  = all_metrics[best_f1_m][1]["F1-score"]
    worst_f1_v = all_metrics[worst_f1_m][1]["F1-score"]
    dl_models  = [m for m in MODELS if m != "xgboost_TL_alignment"]
    dl_f1s     = [all_metrics[m][1]["F1-score"] for m in dl_models]
    dl_min, dl_max = min(dl_f1s), max(dl_f1s)
    return (
        f"觀察：{MODEL_DISPLAY[best_f1_m]} 在二元警報任務表現最佳（F1={best_f1_v:.4f}），"
        f"{MODEL_DISPLAY[worst_f1_m]} 最低（F1={worst_f1_v:.4f}）。"
        f"五個深度學習模型的 F1 介於 {dl_min:.4f}–{dl_max:.4f}，"
        f"差距不大，顯示資料量仍是主要瓶頸。"
    )

def obs_4class():
    best_mf1_m  = max(MODELS, key=lambda m: all_metrics[m][2]["Macro F1"])
    best_wf1_m  = max(MODELS, key=lambda m: all_metrics[m][2]["Weighted F1"])
    best_mf1_v  = all_metrics[best_mf1_m][2]["Macro F1"]
    best_wf1_v  = all_metrics[best_wf1_m][2]["Weighted F1"]
    return (
        f"觀察：{MODEL_DISPLAY[best_mf1_m]} 的 Macro F1 最高（{best_mf1_v:.4f}），"
        f"在少數風險等級的辨識上表現相對穩定。"
        f"{MODEL_DISPLAY[best_wf1_m]} 的 Weighted F1 最高（{best_wf1_v:.4f}），"
        f"顯示其在 high_risk 多數類上預測準確。"
        f"整體而言，mid_risk 與 low_risk 樣本量偏少，所有模型在少數類別上仍有改善空間。"
    )

def obs_ranking(rankings):
    top = rankings[0]
    bottom = rankings[-1]
    return (
        f"綜合三項指標，{top['display']} 排名第一（警報 F1={top['f1']:.4f}，MAE={top['mae']:.2f}）；"
        f"{bottom['display']} 排名末位（警報 F1={bottom['f1']:.4f}，MAE={bottom['mae']:.2f}）。"
    )

# ───────────────────────────────────────────────
# 動態生成限制說明
# ───────────────────────────────────────────────
def build_limitations():
    # 動態取測試筆數（DL 模型取最大）
    dl_totals = [pred_stats[m]["total"] for m in MODELS if m != "xgboost_TL_alignment"]
    dl_total  = max(dl_totals)
    xgb_total = pred_stats["xgboost_TL_alignment"]["total"]

    # inf users（各模型一樣，取 DL 第一個）
    inf_users = pred_stats["bigru"]["inf"]

    return [
        f"所有模型均採 per-user 70/15/15 切分；"
        f"深度學習模型測試集各約 {dl_total} 筆，XGBoost 測試集 {xgb_total} 筆（13 位用戶）。",

        f"各模型測試集中約有 {inf_users} 筆 risk ratio 為 inf，"
        f"原因是這些用戶在訓練期間無 income 交易記錄，導致 monthly_available_cash=0。"
        f"此為資料特性，統一視為最高風險（high_risk），非模型缺陷。",

        "XGBoost + TL 的預訓練資料來源（Kaggle 消費資料）與 GRU/BiLSTM/BiGRU 的預訓練資料"
        "（Walmart 消費資料）不同，遷移效果差異部分來自預訓練資料品質與分佈。",

        "mid_risk 與 low_risk 樣本數偏少，造成所有模型在少數類別上的 Macro F1 偏低，"
        "增加更多用戶資料可望改善此問題。",
    ]

# ───────────────────────────────────────────────
# Word 工具函式
# ───────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # 移除舊的 shd，避免重複
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_header_row(table, headers, bg="1F3864"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)

def add_data_row(table, row_idx, values, bold_col=None,
                 highlight_col=None, highlight_color="FFF2CC"):
    row = table.rows[row_idx]
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = str(v)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.runs[0]
        run.font.size = Pt(9.5)
        if bold_col is not None and i == bold_col:
            run.bold = True
        if highlight_col and i in highlight_col:
            set_cell_bg(cell, highlight_color)
        elif row_idx % 2 == 0:
            set_cell_bg(cell, "F2F2F2")

def best_idx(values, lower_is_better=False):
    try:
        floats = [float(v) for v in values]
        return floats.index(min(floats) if lower_is_better else max(floats))
    except Exception:
        return -1

def add_section_title(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    color = RGBColor(0x1F, 0x38, 0x64) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    p.style.font.color.rgb = color

def add_note(doc, text):
    p   = doc.add_paragraph()
    run = p.add_run("📌 " + text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# ───────────────────────────────────────────────
# 建立文件
# ───────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── 封面 ──
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run("個人財務支出預測模型比較報告")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Machine Learning Model Comparison Report")
r2.font.size = Pt(13)
r2.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f"評估日期：{date.today().isoformat()}").font.size = Pt(11)  # 動態日期

doc.add_page_break()

# ── 1. 研究背景 ──
add_section_title(doc, "一、研究背景與模型說明")

# 動態：用戶數、測試筆數
dl_total = max(pred_stats[m]["total"] for m in MODELS if m != "xgboost_TL_alignment")
dl_users  = len(pred_stats["bigru"]["users"])
doc.add_paragraph(
    f"本報告針對個人財務支出預測任務，比較六個模型的預測表現。"
    f"所有模型以相同的個人交易資料（{dl_users} 位用戶，排除資料量不足的 user4/user5/user6）進行訓練，"
    f"並透過統一的評估框架（output_eval_utils）輸出標準化指標。"
    f"所有模型均採 per-user 70/15/15 切分；"
    f"深度學習模型測試集各 {dl_total} 筆，XGBoost 測試集 {pred_stats['xgboost_TL_alignment']['total']} 筆。"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run("模型架構說明：").bold = True

bullets = [
    ("BiGRU",         "雙向 GRU 循環神經網路，以個人資料直接訓練。"),
    ("BiGRU + TL",    "BiGRU 架構搭配遷移學習（Transfer Learning），以 Walmart 公開消費資料預訓練後 fine-tune 至個人資料，並加入 MMD domain alignment。"),
    ("BiLSTM",        "雙向 LSTM 循環神經網路，以個人資料直接訓練。"),
    ("BiLSTM + TL",   "BiLSTM 架構搭配遷移學習，流程同 BiGRU + TL。"),
    ("GRU + TL",      "GRU 架構搭配遷移學習，並加入 MMD domain alignment 對齊 Walmart 與個人資料分佈。"),
    ("XGBoost + TL",  "XGBoost 梯度提升樹，以 Kaggle 公開消費資料預訓練後遷移至個人資料，並加入 MMD domain alignment。"),
]
for name, desc in bullets:
    p2 = doc.add_paragraph(style="List Bullet")
    p2.add_run(f"{name}：").bold = True
    p2.add_run(desc).font.size = Pt(10)

doc.add_paragraph()
add_note(doc,
    "所有六個模型均使用 per-user 70/15/15 切分（train/valid/test）。"
)

doc.add_page_break()

# ── 2. 迴歸指標 ──
add_section_title(doc, "二、迴歸指標比較（Test 1 — Regression）")
doc.add_paragraph(
    "迴歸指標衡量模型預測每日實際支出金額的準確度（單位：元）。"
    "MAE 為平均絕對誤差、RMSE 為均方根誤差（對大誤差更敏感）、"
    "MAPE 與 SMAPE 為百分比誤差（越低越好）。"
)
doc.add_paragraph()

reg_headers = ["模型", "MAE ↓", "RMSE ↓", "MAPE (%) ↓", "SMAPE (%) ↓"]
reg_table   = doc.add_table(rows=len(MODELS) + 1, cols=5)
reg_table.style     = "Table Grid"
reg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(reg_table, reg_headers)

reg_cols = ["MAE", "RMSE", "MAPE", "SMAPE"]
reg_vals = {k: [all_metrics[m][0][k] for m in MODELS] for k in reg_cols}
best_reg = {k: best_idx(reg_vals[k], lower_is_better=True) for k in reg_cols}

for i, m in enumerate(MODELS):
    reg = all_metrics[m][0]
    row_vals  = [MODEL_DISPLAY[m], f"{reg['MAE']:.2f}", f"{reg['RMSE']:.2f}",
                 f"{reg['MAPE']:.2f}", f"{reg['SMAPE']:.2f}"]
    highlight = [j + 1 for j, k in enumerate(reg_cols) if best_reg[k] == i]
    add_data_row(reg_table, i + 1, row_vals, bold_col=0,
                 highlight_col=highlight, highlight_color="E2EFDA")

add_note(doc, "綠色底色 = 該欄位最佳值（越低越好）")
doc.add_paragraph()
doc.add_paragraph(obs_regression())
doc.add_page_break()

# ── 3. 二元警報指標 ──
add_section_title(doc, "三、二元警報分類指標（Test 2 — Binary Alarm）")
doc.add_paragraph(
    "模型根據預測支出計算 risk ratio，若超過閾值則觸發「財務警報」（alarm=1）。"
    "此為二元分類任務：Precision 衡量警報精準度、Recall 衡量警報覆蓋率、F1 為兩者調和平均。"
)
doc.add_paragraph()

alm_headers = ["模型", "Accuracy ↑", "Precision ↑", "Recall ↑", "F1-score ↑"]
alm_table   = doc.add_table(rows=len(MODELS) + 1, cols=5)
alm_table.style     = "Table Grid"
alm_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(alm_table, alm_headers)

alm_cols = ["Accuracy", "Precision", "Recall", "F1-score"]
alm_vals = {k: [all_metrics[m][1][k] for m in MODELS] for k in alm_cols}
best_alm = {k: best_idx(alm_vals[k]) for k in alm_cols}

for i, m in enumerate(MODELS):
    alm = all_metrics[m][1]
    row_vals  = [MODEL_DISPLAY[m], f"{alm['Accuracy']:.4f}", f"{alm['Precision']:.4f}",
                 f"{alm['Recall']:.4f}", f"{alm['F1-score']:.4f}"]
    highlight = [j + 1 for j, k in enumerate(alm_cols) if best_alm[k] == i]
    add_data_row(alm_table, i + 1, row_vals, bold_col=0,
                 highlight_col=highlight, highlight_color="DDEBF7")

add_note(doc, "藍色底色 = 該欄位最佳值（越高越好）")
doc.add_paragraph()
doc.add_paragraph(obs_alarm())
doc.add_page_break()

# ── 4. 四分類指標 ──
add_section_title(doc, "四、四分類風險等級指標（Test 3 — Risk 4-Class）")
doc.add_paragraph(
    "將 risk ratio 分為四個等級：no_alarm、low_risk、mid_risk、high_risk。"
    "Macro F1 對每個類別等權重平均，反映模型對少數類別的分辨能力；"
    "Weighted F1 依類別樣本數加權，反映整體效能。"
)
doc.add_paragraph()

r4_headers = ["模型", "Accuracy ↑", "Macro F1 ↑", "Weighted F1 ↑", "Macro Prec ↑", "Macro Rec ↑"]
r4_table   = doc.add_table(rows=len(MODELS) + 1, cols=6)
r4_table.style     = "Table Grid"
r4_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(r4_table, r4_headers)

r4_col_keys = ["Accuracy", "Macro F1", "Weighted F1", "Macro Precision", "Macro Recall"]
r4_vals = {k: [all_metrics[m][2][k] for m in MODELS] for k in r4_col_keys}
best_r4 = {k: best_idx(r4_vals[k]) for k in r4_col_keys}

for i, m in enumerate(MODELS):
    r4 = all_metrics[m][2]
    row_vals  = [MODEL_DISPLAY[m], f"{r4['Accuracy']:.4f}", f"{r4['Macro F1']:.4f}",
                 f"{r4['Weighted F1']:.4f}", f"{r4['Macro Precision']:.4f}",
                 f"{r4['Macro Recall']:.4f}"]
    highlight = [j + 1 for j, k in enumerate(r4_col_keys) if best_r4[k] == i]
    add_data_row(r4_table, i + 1, row_vals, bold_col=0,
                 highlight_col=highlight, highlight_color="DDEBF7")

add_note(doc, "藍色底色 = 該欄位最佳值（越高越好）")
doc.add_paragraph()
doc.add_paragraph(obs_4class())
doc.add_page_break()

# ── 5. Confusion Matrix ──
add_section_title(doc, "五、混淆矩陣（四分類）")
doc.add_paragraph(
    "以下列出各模型在四分類任務的混淆矩陣。"
    "列為真實類別（no_alarm / low_risk / mid_risk / high_risk），欄為預測類別。"
    "對角線（綠底）為正確分類數。"
)
doc.add_paragraph()

RISK_LABELS = ["no_alarm", "low_risk", "mid_risk", "high_risk"]

for m in MODELS:
    cm = all_metrics[m][2]["Confusion Matrix"]
    add_section_title(doc, MODEL_DISPLAY[m], level=2)

    cm_table = doc.add_table(rows=5, cols=5)
    cm_table.style     = "Table Grid"
    cm_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header_row = cm_table.rows[0]
    header_row.cells[0].text = "真實 \\ 預測"
    set_cell_bg(header_row.cells[0], "1F3864")
    r0 = header_row.cells[0].paragraphs[0].runs[0]
    r0.bold = True; r0.font.color.rgb = RGBColor(255,255,255); r0.font.size = Pt(9)

    for j, lbl in enumerate(RISK_LABELS):
        cell = header_row.cells[j + 1]
        cell.text = lbl
        set_cell_bg(cell, "2E74B5")
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.runs[0]
        run.bold = True; run.font.color.rgb = RGBColor(255,255,255); run.font.size = Pt(9)

    for row_i, row_label in enumerate(RISK_LABELS):
        row = cm_table.rows[row_i + 1]
        row.cells[0].text = row_label
        set_cell_bg(row.cells[0], "BDD7EE")
        lp = row.cells[0].paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.runs[0]; lr.bold = True; lr.font.size = Pt(9)

        for col_j, val in enumerate(cm[row_i]):
            cell = row.cells[col_j + 1]
            cell.text = str(val)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run  = para.runs[0]
            run.font.size = Pt(9)
            if col_j == row_i:          # 對角線
                set_cell_bg(cell, "C6EFCE")
                run.bold = True

    doc.add_paragraph()

doc.add_page_break()

# ── 6. 綜合排名（動態）──
add_section_title(doc, "六、綜合排名與建議")
doc.add_paragraph(
    "以下依「二元警報 F1（主排序）→ MAE（次排序）」動態計算排名，"
    "綜合回歸誤差、警報分類、四分類三個面向給出整體評估。"
)
doc.add_paragraph()

rank_headers = ["排名", "模型", "MAE ↓", "警報 F1 ↑", "4-Class Macro F1 ↑"]
rank_table   = doc.add_table(rows=len(rankings) + 1, cols=5)
rank_table.style     = "Table Grid"
rank_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_header_row(rank_table, rank_headers)

rank_colors = ["FFF2CC", "E2EFDA", "E2EFDA", "FFFFFF", "FFFFFF", "F2F2F2"]
for i, r in enumerate(rankings):
    row = rank_table.rows[i + 1]
    vals = [str(r["rank"]), r["display"],
            f"{r['mae']:.2f}", f"{r['f1']:.4f}", f"{r['macro_f1']:.4f}"]
    bg = rank_colors[i] if i < len(rank_colors) else "FFFFFF"
    for j, v in enumerate(vals):
        cell = row.cells[j]
        cell.text = v
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = para.runs[0]
        run.font.size = Pt(9.5)
        if j == 0:
            run.bold = True
    set_cell_bg(row.cells[0], bg)

doc.add_paragraph()
doc.add_paragraph(obs_ranking(rankings))

# ── 7. 資料說明與限制（動態）──
doc.add_paragraph()
add_section_title(doc, "七、資料說明與限制", level=2)
doc.add_paragraph()
for lim in build_limitations():
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(lim).font.size = Pt(10)

# ── 頁尾 ──
for section in doc.sections:
    footer = section.footer
    fp     = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("個人財務支出預測模型比較報告  |  Financial Agent Project").font.size = Pt(8)

doc.save(OUTPUT_PATH)
print(f"✅ 報告已儲存：{OUTPUT_PATH}")
